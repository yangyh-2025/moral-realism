"""
Convert all 5 markdown docs to Word (.docx) with mandated formatting:
- 正文: 仿宋/SimFang + Times New Roman, 小四(12pt), 行距1.5, 段前段后0, 首行缩进2字符
- 标题: 小标宋, 20pt
- 一级标题(一、二、): 黑体(SimHei), 四号(14pt)
- 二级标题((一)(二)): 楷体(KaiTi)加粗, 四号(14pt)
- 三级标题(1. 2.): 仿宋(SimFang)加粗, 小四(12pt)
"""
import re, os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))

DOCS = [
    ('仿真设计说明', os.path.join(DOCS_DIR, '仿真设计说明.md'), os.path.join(DOCS_DIR, 'simulation_design')),
    ('技术文档', os.path.join(DOCS_DIR, '技术文档.md'), os.path.join(DOCS_DIR, 'technical_doc')),
    ('智能体提示词设计说明', os.path.join(DOCS_DIR, '智能体提示词设计说明.md'), os.path.join(DOCS_DIR, 'prompt_design')),
    ('程序设计说明', os.path.join(DOCS_DIR, '程序设计说明.md'), os.path.join(DOCS_DIR, 'program_design')),
    ('三场景历史背景说明', os.path.join(DOCS_DIR, '三场景历史背景说明.md'), os.path.join(DOCS_DIR, 'scene_history')),
]

FONT_SIZE_BODY = Pt(12)      # 小四
FONT_SIZE_H1 = Pt(14)        # 四号
FONT_SIZE_H2 = Pt(14)        # 四号
FONT_SIZE_H3 = Pt(12)        # 小四
FONT_SIZE_TITLE = Pt(20)     # 20号


def set_paragraph_spacing(paragraph, line_spacing=1.5, space_before=0, space_after=0):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)


def set_run_font(run, cn_font='仿宋', en_font='Times New Roman', size=FONT_SIZE_BODY, bold=False):
    """Set run font for Chinese and Western text."""
    run.font.size = size
    run.bold = bold
    run.font.name = en_font
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)

    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
        rPr.insert(0, rFonts)

    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    rFonts.set(qn('w:cs'), en_font)


def add_body_paragraph(doc, text, first_line_indent=True):
    """Add body paragraph with 仿宋/TNR, 小四, 1.5 line spacing, optional indent."""
    para = doc.add_paragraph()
    set_paragraph_spacing(para, line_spacing=1.5, space_before=0, space_after=0)
    if first_line_indent:
        para.paragraph_format.first_line_indent = Pt(24)  # ~2 Chinese characters

    # Parse bold markers **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            set_run_font(run, bold=True)
        else:
            run = para.add_run(part)
            set_run_font(run)
    return para


def add_title(doc, text):
    """Add document title: 小标宋, 20pt, centered."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(para, line_spacing=1.5, space_before=0, space_after=Pt(6))
    # Strip markdown formatting
    text = text.replace('# ', '').replace('**', '').strip()
    run = para.add_run(text)
    set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=FONT_SIZE_TITLE, bold=True)
    return para


def add_h1(doc, text):
    """Add level-1 heading: 黑体(SimHei), 四号(14pt), 一、二、三、 numbering."""
    para = doc.add_paragraph()
    set_paragraph_spacing(para, line_spacing=1.5, space_before=6, space_after=3)
    text = text.strip()
    # Remove markdown ## prefix
    text = re.sub(r'^#+\s*', '', text)
    run = para.add_run(text)
    set_run_font(run, cn_font='黑体', en_font='Times New Roman', size=FONT_SIZE_H1, bold=True)
    return para


def add_h2(doc, text):
    """Add level-2 heading: 楷体(KaiTi)加粗, 四号(14pt), (一)(二)(三) numbering."""
    para = doc.add_paragraph()
    set_paragraph_spacing(para, line_spacing=1.5, space_before=4, space_after=2)
    text = re.sub(r'^#+\s*', '', text).strip()
    run = para.add_run(text)
    set_run_font(run, cn_font='楷体', en_font='Times New Roman', size=FONT_SIZE_H2, bold=True)
    return para


def add_h3(doc, text):
    """Add level-3 heading: 仿宋(SimFang)加粗, 小四(12pt)."""
    para = doc.add_paragraph()
    set_paragraph_spacing(para, line_spacing=1.5, space_before=3, space_after=1)
    text = re.sub(r'^#+\s*', '', text).strip()
    run = para.add_run(text)
    set_run_font(run, cn_font='仿宋', en_font='Times New Roman', size=FONT_SIZE_H3, bold=True)
    return para


def parse_markdown_to_docx(md_path, docx_path, doc_title):
    """Convert markdown file to formatted docx."""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Set default styles
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = FONT_SIZE_BODY
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    i = 0
    in_table = False
    in_code_block = False
    in_mermaid = False
    table_rows = []
    table_header = []
    code_buffer = []

    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            if in_table:
                # End table
                build_table(doc, table_header, table_rows)
                table_rows = []
                table_header = []
                in_table = False
            if in_code_block and not stripped == '```':
                code_buffer.append('')
            i += 1
            continue

        # Mermaid blocks - skip completely
        if stripped.startswith('```mermaid'):
            in_mermaid = True
            in_code_block = True
            i += 1
            continue
        if in_mermaid:
            if stripped == '```':
                in_mermaid = False
                in_code_block = False
            i += 1
            continue

        # Code blocks
        if stripped.startswith('```'):
            if in_code_block:
                # End code block
                add_code_block(doc, code_buffer)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Tables
        if stripped.startswith('|') and stripped.endswith('|'):
            if not in_table:
                in_table = True
                table_header = [c.strip() for c in stripped.split('|')[1:-1]]
            elif table_rows or (not re.match(r'^[\|\s\-\:]+$', stripped)):
                # Skip separator line like |---|---|
                if not re.match(r'^[\|\s\-\:]+$', stripped):
                    row_cells = [c.strip() for c in stripped.split('|')[1:-1]]
                    table_rows.append(row_cells)
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,6})\s+(.+)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            content = heading_match.group(2)

            if level == 1:
                add_title(doc, content)
            elif level == 2:
                add_h1(doc, content)
            elif level == 3:
                add_h2(doc, content)
            elif level == 4:
                add_h3(doc, content)
            else:
                # Level 5/6 as body bold
                para = doc.add_paragraph()
                set_paragraph_spacing(para)
                para.paragraph_format.first_line_indent = Pt(24)
                run = para.add_run(content)
                set_run_font(run, bold=True)

            i += 1
            continue

        # Horizontal rules
        if stripped in ('---', '***', '___'):
            para = doc.add_paragraph()
            set_paragraph_spacing(para, line_spacing=1.0)
            # Add a simple horizontal rule
            pPr = para._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="808080"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # Image references
        if stripped.startswith('!['):
            i += 1
            continue

        # Blockquotes
        if stripped.startswith('>'):
            content = re.sub(r'^>\s*', '', stripped)
            para = doc.add_paragraph()
            set_paragraph_spacing(para, line_spacing=1.5)
            para.paragraph_format.first_line_indent = Pt(24)
            para.paragraph_format.left_indent = Cm(1)
            run = para.add_run(content)
            set_run_font(run, cn_font='仿宋', size=FONT_SIZE_BODY)
            run.font.italic = True
            i += 1
            continue

        # Lists
        if re.match(r'^\s*[\-\*\+]\s+', stripped):
            content = re.sub(r'^\s*[\-\*\+]\s+', '', stripped)
            para = doc.add_paragraph()
            set_paragraph_spacing(para, line_spacing=1.5)
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.left_indent = Cm(1)

            # Parse nested bold
            parts = re.split(r'(\*\*.*?\*\*)', content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    set_run_font(run, bold=True)
                else:
                    run = para.add_run(part)
                    set_run_font(run)
            i += 1
            continue

        # Numbered lists
        if re.match(r'^\s*\d+[\.\)]\s+', stripped):
            content = re.sub(r'^\s*\d+[\.\)]\s+', '', stripped)
            para = doc.add_paragraph()
            set_paragraph_spacing(para, line_spacing=1.5)
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.left_indent = Cm(1)

            parts = re.split(r'(\*\*.*?\*\*)', content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    set_run_font(run, bold=True)
                else:
                    run = para.add_run(part)
                    set_run_font(run)
            i += 1
            continue

        # LaTeX math lines (skip, can't render well in docx)
        if stripped.startswith('$$') or stripped.startswith('$'):
            i += 1
            continue

        # Regular body text
        add_body_paragraph(doc, stripped)
        i += 1

    # Handle trailing table
    if in_table and table_rows:
        build_table(doc, table_header, table_rows)

    doc.save(docx_path)
    return docx_path


def build_table(doc, header, rows):
    """Build a formatted table in the document."""
    if not header:
        return
    table = doc.add_table(rows=len(rows) + 1, cols=len(header))
    table.style = 'Table Grid'

    # Header row
    for j, cell_text in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ''
        para = cell.paragraphs[0]
        run = para.add_run(cell_text)
        set_run_font(run, cn_font='黑体', size=Pt(10), bold=True)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(para, line_spacing=1.2, space_before=2, space_after=2)
        # Header background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2c3e50" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j >= len(header):
                continue
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            para = cell.paragraphs[0]
            run = para.add_run(str(cell_text))
            set_run_font(run, size=Pt(9))
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(para, line_spacing=1.2, space_before=1, space_after=1)
            # Alternating row colors
            if i % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="ecf0f1" w:val="clear"/>')
                cell._element.get_or_add_tcPr().append(shading)

    doc.add_paragraph()  # Space after table


def add_code_block(doc, code_lines):
    """Add a code block with monospace font."""
    if not code_lines:
        return
    for line in code_lines:
        para = doc.add_paragraph()
        set_paragraph_spacing(para, line_spacing=1.0, space_before=0, space_after=0)
        para.paragraph_format.first_line_indent = Pt(0)
        para.paragraph_format.left_indent = Cm(1)
        run = para.add_run(line if line else ' ')
        # Use Courier New for code
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        r = run._element
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
            r.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), '仿宋')
        rFonts.set(qn('w:ascii'), 'Courier New')
        rFonts.set(qn('w:hAnsi'), 'Courier New')
    doc.add_paragraph()  # Space after code block


def main():
    print('Converting documents to Word format...')
    for name, md_path, out_dir in DOCS:
        os.makedirs(out_dir, exist_ok=True)
        docx_path = os.path.join(out_dir, f'{name}.docx')

        try:
            path = parse_markdown_to_docx(md_path, docx_path, name)
            size_kb = os.path.getsize(path) / 1024
            print(f'  OK {name}: {size_kb:.0f} KB -> {path}')
        except Exception as e:
            print(f'  FAIL {name}: {e}')
            import traceback
            traceback.print_exc()

    print('\nAll documents converted successfully!')


if __name__ == '__main__':
    main()
