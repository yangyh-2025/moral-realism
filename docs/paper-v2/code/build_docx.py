# -*- coding: utf-8 -*-
"""Build Chinese academic DOCX — simplified: python-docx + ZIP footnote fix."""
import re, os, zipfile, shutil
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MD_PATH = os.path.join(BASE, 'paper', '论文_实验部分.md')
OUT_PATH = os.path.join(BASE, 'paper', '论文_实验部分.docx')
FIG_DIR = os.path.join(BASE, 'figures')
NS_W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

def _font(run, cn, en='Times New Roman', size=12, bold=False):
    run.font.size = Pt(size); run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is not None: rPr.remove(rf)
    rf = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{cn}" w:ascii="{en}" w:hAnsi="{en}" w:cs="{en}"/>')
    rPr.insert(0, rf)

def _clean(text):
    text = text.replace('\\"', '"').replace('\\*', '*').replace('\\[', '[').replace('\\]', ']')
    text = re.sub(r'-{3,}', '——', text)
    text = re.sub(r'\$\^\{([^}]+)\}\$', r'\1', text)
    text = re.sub(r'~([^~]+)~', r'\1', text)
    text = re.sub(r'\^([^^]+)\^', r'\1', text)
    text = text.replace('...', '……')
    return text

# ---- footnotes ----
class Footnotes:
    def __init__(self, doc):
        self.doc = doc; self.n = 1
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        for rel in doc.part.rels.values():
            if 'footnotes' in rel.reltype:
                self.p = rel.target_part; return
        xml = ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:footnote w:type="separator" w:id="-1"><w:p><w:pPr><w:spacing w:line="240" w:lineRule="auto"/></w:pPr>'
            '<w:r><w:separator/></w:r></w:p></w:footnote>'
            '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:pPr><w:spacing w:line="240" w:lineRule="auto"/></w:pPr>'
            '<w:r><w:continuationSeparator/></w:r></w:p></w:footnote></w:footnotes>')
        ct = 'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml'
        self.p = Part(PackURI('/word/footnotes.xml'), ct, xml.encode('utf-8'), doc.part.package)
        doc.part.relate_to(self.p, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes')

    def add(self, text):
        fid = self.n; self.n += 1
        tree = etree.fromstring(self.p.blob)
        el = parse_xml(
            f'<w:footnote {nsdecls("w")} w:id="{fid}"><w:p>'
            f'<w:pPr><w:spacing w:line="240" w:lineRule="auto"/><w:jc w:val="both"/></w:pPr>'
            f'<w:r><w:rPr><w:rStyle w:val="FootnoteReference"/>'
            f'<w:rFonts w:eastAsia="FangSong" w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>'
            f'<w:sz w:val="21"/><w:szCs w:val="21"/></w:rPr><w:footnoteRef/></w:r>'
            f'<w:r><w:rPr><w:rFonts w:eastAsia="FangSong" w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>'
            f'<w:sz w:val="21"/><w:szCs w:val="21"/></w:rPr><w:t xml:space="preserve"> {text}</w:t></w:r>'
            f'</w:p></w:footnote>')
        tree.append(el)
        self.p._blob = etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)
        return fid

    def ref(self, para, fid):
        para._element.append(parse_xml(
            f'<w:r {nsdecls("w")}><w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr>'
            f'<w:footnoteReference w:id="{fid}"/></w:r>'))

# ---- parse ----
def parse_md(path):
    with open(path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    blocks = []; first_h1 = True; i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line: i += 1; continue
        # Title (first #)
        if (m := re.match(r'^# (.+)$', line)) and first_h1:
            blocks.append(('title', _clean(m.group(1)))); first_h1 = False; i += 1; continue
        # H2
        if m := re.match(r'^## (.+)$', line):
            blocks.append(('h2', _clean(m.group(1)))); i += 1; continue
        # H3
        if m := re.match(r'^### (.+)$', line):
            blocks.append(('h3', _clean(m.group(1)))); i += 1; continue
        # Image
        if m := re.match(r'^!\[(.+)\]\((.+)\)$', line):
            blocks.append(('image', m.group(1), m.group(2))); i += 1; continue
        # Standalone **term**
        if (m := re.match(r'^\*\*(.+)\*\*$', line)) and not re.search(r'[。，；：]', m.group(1)):
            blocks.append(('term', _clean(m.group(1)))); i += 1; continue
        # Table
        if re.match(r'^\|', line):
            rows = []
            while i < len(lines) and re.match(r'^\|', lines[i].rstrip()):
                row = lines[i].rstrip()
                if not re.match(r'^[\|\s\-:]+$', row): rows.append(row)
                i += 1
            if rows: blocks.append(('table', rows))
            continue
        # Display math $$
        if line.strip() == '$$':
            ml = []; i += 1
            while i < len(lines) and lines[i].strip() != '$$':
                ml.append(lines[i].rstrip()); i += 1
            if i < len(lines): i += 1
            if ml: blocks.append(('math', '\n'.join(ml)))
            continue
        # Body
        bl = [line]; i += 1
        while i < len(lines) and lines[i].strip():
            bl.append(lines[i].rstrip()); i += 1
        blocks.append(('body', ' '.join(bl)))
    return blocks

# ---- inline render ----
def _render_inline(p, text):
    text = _clean(text)
    text = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'\1', text)
    text = re.sub(r'\[([^\]]*)\]\([^)]+\)', r'\1', text)
    text = re.sub(r'\[([^\]]*)\]\[[^\]]*\]', r'\1', text)
    # **bold**, $math$
    pattern = r'(\*\*(.+?)\*\*|\$([^$]+)\$)'
    last = 0
    for m in re.finditer(pattern, text):
        if m.start() > last:
            r = p.add_run(text[last:m.start()]); _font(r, 'FangSong')
        if m.group(2):
            r = p.add_run(m.group(2)); _font(r, 'FangSong', bold=True)
        elif m.group(3):
            math_clean = re.sub(r'[_^]\{([^}]+)\}', r'\1', m.group(3))
            r = p.add_run(math_clean); _font(r, 'FangSong')
        last = m.end()
    if last < len(text):
        r = p.add_run(text[last:]); _font(r, 'FangSong')

# ---- table render ----
def _render_table(doc, rows_data):
    ncols = max(len([c.strip() for c in r.split('|')[1:-1]]) for r in rows_data)
    nrows = len(rows_data)
    usable = 9072; col_w = usable // ncols
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = 'Table Grid'
    for ri, row in enumerate(rows_data):
        cells = [c.strip() for c in row.split('|')[1:-1]]
        for ci, ct in enumerate(cells):
            if ci >= ncols: break
            cell = table.cell(ri, ci)
            cell.paragraphs[0].clear()
            ct = re.sub(r'\*\*(.+?)\*\*', r'\1', ct)
            r = cell.paragraphs[0].add_run(_clean(ct))
            sz = 10.5 if (ri == 0 or ncols >= 5) else 12
            _font(r, 'FangSong', size=sz, bold=(ri == 0))
            for ptag in cell.paragraphs:
                ptag.paragraph_format.space_before = Pt(1)
                ptag.paragraph_format.space_after = Pt(1)
                ptag.paragraph_format.line_spacing = 1.2
    doc.add_paragraph()

# ---- main ----
def main():
    doc = Document()
    ns_style = doc.styles['Normal']
    for tag in (qn('w:pPr'), qn('w:rPr')):
        el = ns_style.element.find(tag)
        if el is not None: ns_style.element.remove(el)

    # Section setup — margins + page number via native footer API
    for sec in doc.sections:
        sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.17); sec.right_margin = Cm(3.17)
        # Native page number footer
        footer = sec.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run()
        _font(fr, 'Times New Roman', size=9)
        fld = parse_xml(
            f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        fr._element.append(fld)
        itxt = parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        fr._element.append(itxt)
        fe = parse_xml(
            f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        fr._element.append(fe)

    blocks = parse_md(MD_PATH)
    fn_obj = Footnotes(doc)

    # Collect footnote defs — scan raw file line by line
    with open(MD_PATH, 'r', encoding='utf-8') as f:
        raw_text = f.read()
    fn_defs = {}
    for m in re.finditer(r'\[\^(\d+)\]:\s*(.*?)\s*(?=\[\^|\Z)', raw_text, re.DOTALL):
        body = m.group(2).strip()
        # Clean up trailing artifacts
        body = re.sub(r'\n\s*$', '', body)
        fn_defs[m.group(1)] = body
    print(f'  Footnote definitions found: {sorted(fn_defs.keys())}')
    # Remove footnote definition paragraphs
    fn_pat = re.compile(r'^\[\^(\d+)\]')
    blocks = [b for b in blocks if not (b[0] == 'body' and fn_pat.match(b[1]))]

    for blk in blocks:
        t = blk[0]
        if t == 'title':
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.space_before = Pt(0); pf.space_after = Pt(12); pf.line_spacing = 1.5
            r = p.add_run(blk[1]); _font(r, 'STXinwei', size=20)

        elif t == 'h2':
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.space_before = Pt(12); pf.space_after = Pt(6); pf.line_spacing = 1.5
            r = p.add_run(blk[1]); _font(r, 'SimHei', size=14, bold=True)

        elif t == 'h3':
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.space_before = Pt(6); pf.space_after = Pt(3); pf.line_spacing = 1.5
            r = p.add_run(blk[1]); _font(r, 'KaiTi', size=14, bold=True)

        elif t == 'term':
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.space_before = Pt(6); pf.space_after = Pt(0); pf.line_spacing = 1.5
            r = p.add_run(blk[1]); _font(r, 'FangSong', size=12, bold=True)

        elif t == 'image':
            alt, rel_path = blk[1], blk[2]
            fname = os.path.basename(rel_path)
            img_path = os.path.join(FIG_DIR, fname)
            if os.path.exists(img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(3)
                r = p.add_run()
                try:
                    r.add_picture(img_path, width=Inches(5.5))
                except Exception as e:
                    print(f'  Image failed: {fname}: {e}')
                cp = doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.paragraph_format.space_after = Pt(8)
                _font(cp.add_run(alt), 'KaiTi', size=10)
            else:
                print(f'  Image not found: {img_path}')

        elif t == 'math':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(blk[1]); _font(r, 'FangSong', bold=False)

        elif t == 'table':
            _render_table(doc, blk[1])

        elif t == 'body':
            text = blk[1]
            fn_refs = list(re.finditer(r'\[\^(\d+)\]', text))
            if fn_refs:
                last_end = 0
                p = doc.add_paragraph()
                pf = p.paragraph_format
                pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                pf.first_line_indent = Pt(24)
                pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1.5

                for fm in fn_refs:
                    if fm.start() > last_end:
                        _render_inline(p, text[last_end:fm.start()])
                    fn_num = fm.group(1)
                    fn_body = fn_defs.get(fn_num, '')
                    if fn_body:
                        fid = fn_obj.add(fn_body)
                        fn_obj.ref(p, fid)
                    last_end = fm.end()
                if last_end < len(text):
                    _render_inline(p, text[last_end:])
            else:
                p = doc.add_paragraph()
                pf = p.paragraph_format
                pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                pf.first_line_indent = Pt(24)
                pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1.5
                _render_inline(p, text)

    doc.save(OUT_PATH)

    # ── Post-process: footnote ①②③ per-page + sz fixes ──
    tmp = OUT_PATH + '.tmp'
    with zipfile.ZipFile(OUT_PATH, 'r') as zin:
        with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                fname = item.filename

                if fname == 'word/settings.xml':
                    tree = etree.fromstring(data)
                    fnPr = tree.find(f'{NS_W}footnotePr')
                    if fnPr is None:
                        fnPr = parse_xml(f'<w:footnotePr {nsdecls("w")} />')
                        tree.append(fnPr)
                    fnPr.append(parse_xml(f'<w:numFmt {nsdecls("w")} w:val="ideographEnclosedCircle"/>'))
                    fnPr.append(parse_xml(f'<w:numRestart {nsdecls("w")} w:val="newPage"/>'))
                    data = etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)

                elif fname == 'word/styles.xml':
                    tree = etree.fromstring(data)
                    # DocDefaults sz=21
                    dd = tree.find(f'{NS_W}docDefaults')
                    if dd is None:
                        dd = parse_xml(f'<w:docDefaults {nsdecls("w")} />')
                        tree.insert(0, dd)
                    rpr_def = dd.find(f'{NS_W}rPrDefault')
                    if rpr_def is None:
                        rpr_def = parse_xml(f'<w:rPrDefault {nsdecls("w")} />')
                        dd.append(rpr_def)
                    rpr_dd = rpr_def.find(f'{NS_W}rPr')
                    if rpr_dd is None:
                        rpr_dd = parse_xml(f'<w:rPr {nsdecls("w")} />')
                        rpr_def.append(rpr_dd)
                    for tagn in ['sz', 'szCs']:
                        ex = rpr_dd.find(f'{NS_W}{tagn}')
                        if ex is not None: rpr_dd.remove(ex)
                    rpr_dd.append(parse_xml(f'<w:sz {nsdecls("w")} w:val="21"/>'))
                    rpr_dd.append(parse_xml(f'<w:szCs {nsdecls("w")} w:val="21"/>'))

                    # Fix FootnoteText + FootnoteReference styles
                    for sid in ['FootnoteText', 'FootnoteReference']:
                        for st in tree.findall(f'{NS_W}style'):
                            if st.get(f'{NS_W}styleId') == sid:
                                st_rpr = st.find(f'{NS_W}rPr')
                                if st_rpr is None:
                                    st_rpr = parse_xml(f'<w:rPr {nsdecls("w")} />')
                                    st.append(st_rpr)
                                for tagn in ['sz', 'szCs']:
                                    ex = st_rpr.find(f'{NS_W}{tagn}')
                                    if ex is not None: st_rpr.remove(ex)
                                st_rpr.append(parse_xml(f'<w:sz {nsdecls("w")} w:val="21"/>'))
                                st_rpr.append(parse_xml(f'<w:szCs {nsdecls("w")} w:val="21"/>'))
                                break
                    data = etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)

                zout.writestr(item, data)
    shutil.move(tmp, OUT_PATH)

    print(f'Done! -> {OUT_PATH}')
    print(f'  Footnotes: {fn_obj.n - 1}')
    n_img = sum(1 for b in blocks if b[0]=='image')
    print(f'  Images: {n_img}')
    n_tbl = sum(1 for b in blocks if b[0]=='table')
    print(f'  Tables: {n_tbl}')

if __name__ == '__main__':
    main()
