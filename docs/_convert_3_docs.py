# -*- coding: utf-8 -*-
"""Convert 仿真设计说明, 模型校验章节, 程序设计说明 to DOCX."""
import re, os, sys
sys.path.insert(0, '.')
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

TARGETS = [
    ('docs/simulation_design/仿真设计说明.md', 'docs/simulation_design/仿真设计说明.docx', 'docs/simulation_design/figures'),
    ('docs/model_validation/模型校验章节.md', 'docs/model_validation/模型校验章节.docx', 'docs/model_validation/figures'),
    ('docs/program_design/程序设计说明.md', 'docs/program_design/程序设计说明.docx', 'docs/program_design/figures'),
]

def _font(run, cn, en='Times New Roman', size=12, bold=False, italic=False):
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rPr.insert(0, rf)
    for k, v in [('eastAsia', cn), ('ascii', en), ('hAnsi', en), ('cs', en)]:
        rf.set(qn(f'w:{k}'), v)

def _spacing(para, before=0, after=0, line=1.5):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line

def _indent(para, pt=24):
    para.paragraph_format.first_line_indent = Pt(pt)

def _add_rich(para, text, cn='FangSong', en='Times New Roman', size=12, bold=False):
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*)'
    last = 0
    for m in re.finditer(pattern, text):
        if m.start() > last:
            r = para.add_run(text[last:m.start()])
            _font(r, cn, en, size, bold)
        if m.group(2):
            r = para.add_run(m.group(2))
            _font(r, cn, en, size, True)
        elif m.group(3):
            r = para.add_run(m.group(3))
            _font(r, cn, en, size, bold, True)
        last = m.end()
    if last < len(text):
        r = para.add_run(text[last:])
        _font(r, cn, en, size, bold)

def _is_table_sep(line):
    return bool(re.match(r'^\|[\s\-:|]+\|$', line.strip()))

def _is_table_row(line):
    s = line.strip()
    return s.startswith('|') and s.endswith('|')

def _peek_table(lines, i):
    tbl = [lines[i]]
    j = i + 1
    while j < len(lines):
        s = lines[j].strip()
        if _is_table_row(s) or _is_table_sep(s):
            tbl.append(lines[j])
            j += 1
        else:
            break
    return tbl, j

def _add_table(doc, lines):
    clean = [l for l in lines if not _is_table_sep(l)]
    if len(clean) < 2:
        return
    header = [c.strip() for c in clean[0].split('|') if c.strip()]
    data = []
    for l in clean[1:]:
        cells = [c.strip() for c in l.split('|') if c.strip()]
        if cells:
            data.append(cells)
    if not header:
        return
    ncols = len(header)
    nrows = 1 + len(data)
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = 'Table Grid'
    sz = 12 if ncols <= 5 else 9
    for ci, ct in enumerate(header):
        cell = table.cell(0, ci)
        cell.paragraphs[0].clear()
        _add_rich(cell.paragraphs[0], ct, 'SimHei', 'Times New Roman', sz)
        _spacing(cell.paragraphs[0], 0, 0, 1.1)
    for ri, row in enumerate(data):
        for ci, ct in enumerate(row):
            if ci >= ncols:
                continue
            cell = table.cell(ri + 1, ci)
            cell.paragraphs[0].clear()
            _add_rich(cell.paragraphs[0], ct, 'FangSong', 'Times New Roman', sz)
            _spacing(cell.paragraphs[0], 0, 0, 1.1)
    p = doc.add_paragraph()
    _spacing(p, 6, 0, 1.5)

def _is_image(line):
    return bool(re.match(r'^\!\[(.+?)\]\(.+\)$', line.strip()))

def _image_alt(line):
    m = re.match(r'^\!\[(.+?)\]\(.+\)$', line.strip())
    return m.group(1) if m else None

def _image_name(alt):
    """Map alt text to filename."""
    mapping = {
        '模型校验三场景时间线': 'fig4_three_scenes_timeline.png',
        'CINC综合国力指数计算架构': 'fig5_cinc_calculation.png',
        '四种领导类型行为偏好对比': 'fig3_leader_types_radar.png',
        '20项标准行为空间配置总览': 'fig2_action_space.png',
        '仿真单轮执行流程': 'fig1_simulation_flow.png',
        'Figure 1': 'fig1_overall_comparison.png',
        'Figure 2': 'fig2_per_round_timeseries.png',
        'Figure 3': 'fig3_error_structure.png',
        'Figure 4': 'fig4_f1_distribution.png',
        'Figure 5': 'fig5_cross_country_f1.png',
        'Figure 6': 'fig6_country_stability.png',
        'Figure 7': 'fig7_bloc_alignment.png',
        'Figure 8': 'fig8_italy_case_study.png',
        'Figure 9': 'fig9_gp_independence.png',
        'Figure 10': 'fig10_summary_dashboard.png',
        '核心仿真引擎模块架构': 'fig1_module_architecture.png',
        '数据持久化实体关系概览': 'fig2_data_er_diagram.png',
        '前端研究分析仪表板': 'fig3_frontend_dashboard.png',
        '前端并发控制机制': 'fig4_concurrency_control.png',
    }
    return mapping.get(alt, alt)

def convert(md_path, out_path, fig_dir):
    print(f'Converting: {md_path}')
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    lines = content.split('\n')

    doc = Document()
    ns_el = doc.styles['Normal']
    for tag in (qn('w:pPr'), qn('w:rPr')):
        el = ns_el.element.find(tag)
        if el is not None:
            ns_el.element.remove(el)

    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.18)
        sec.right_margin = Cm(3.18)

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line == '---':
            i += 1
            continue

        # Image
        if _is_image(line):
            alt = _image_alt(line)
            fname = _image_name(alt)
            fpath = os.path.join(fig_dir, fname)
            # try fuzzy match
            if not os.path.exists(fpath) and os.path.isdir(fig_dir):
                for f in os.listdir(fig_dir):
                    base = os.path.splitext(f)[0]
                    if alt in base or base in alt:
                        fpath = os.path.join(fig_dir, f)
                        break
            if os.path.exists(fpath):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _spacing(p, 3, 3, 1.5)
                r = p.add_run()
                r.add_picture(fpath, width=Inches(5.5))
            # image caption
            if i + 1 < len(lines) and lines[i + 1].strip().startswith('*图：'):
                cap = re.sub(r'^\*(.+)\*$', r'\1', lines[i + 1].strip())
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _spacing(p, 0, 3, 1.5)
                _add_rich(p, cap, 'FangSong', 'Times New Roman', 12)
                i += 1
            i += 1
            continue

        # Table
        if _is_table_row(line):
            tbl, next_i = _peek_table(lines, i)
            print(f'  Table: {len(tbl)} rows')
            _add_table(doc, tbl)
            i = next_i
            continue

        # Title (# )
        m1 = re.match(r'^# (.+)$', line)
        if m1 and not line.startswith('## '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _spacing(p, 12, 12, 2.0)
            _add_rich(p, m1.group(1), 'STXinwei', 'Times New Roman', 22)
            i += 1
            continue

        # H1 (## )
        m2 = re.match(r'^## (.+)$', line)
        if m2:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _spacing(p, 6, 0, 1.5)
            _add_rich(p, m2.group(1), 'SimHei', 'Times New Roman', 15, True)
            i += 1
            continue

        # H2 (### )
        m3 = re.match(r'^### (.+)$', line)
        if m3:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _spacing(p, 4, 0, 1.5)
            _add_rich(p, m3.group(1), 'KaiTi', 'Times New Roman', 14, True)
            i += 1
            continue

        # H3 (#### )
        m4 = re.match(r'^#### (.+)$', line)
        if m4:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _spacing(p, 3, 0, 1.5)
            _add_rich(p, m4.group(1), 'FangSong', 'Times New Roman', 12, True)
            i += 1
            continue

        # Math block ($$...$$)
        if line.startswith('$$') and line.endswith('$$'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _spacing(p, 3, 3, 1.5)
            _add_rich(p, line.strip('$'), 'Consolas', 'Consolas', 10)
            i += 1
            continue

        # Inline math
        if '$$' in line:
            parts = re.split(r'(\$\$.*?\$\$)', line)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _spacing(p, 0, 0, 1.5)
            _indent(p, 24)
            for part in parts:
                if part.startswith('$$') and part.endswith('$$'):
                    r = p.add_run(part.strip('$'))
                    _font(r, 'Consolas', 'Consolas', 10)
                else:
                    _add_rich(p, part, 'FangSong', 'Times New Roman', 12)
            i += 1
            continue

        # Bullet
        if re.match(r'^[-*]\s+', line):
            text = re.sub(r'^[-*]\s+', '', line)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _spacing(p, 0, 0, 1.5)
            pf = p.paragraph_format
            pf.left_indent = Cm(1)
            pf.first_line_indent = Cm(-0.5)
            r = p.add_run('  ')
            _font(r, 'FangSong', 'Times New Roman', 12)
            _add_rich(p, text, 'FangSong', 'Times New Roman', 12)
            i += 1
            continue

        # Blockquote
        if line.startswith('> '):
            text = line[2:]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _spacing(p, 0, 0, 1.5)
            pf = p.paragraph_format
            pf.left_indent = Cm(1.5)
            _add_rich(p, text, 'KaiTi', 'Times New Roman', 12)
            i += 1
            continue

        # Body
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _spacing(p, 0, 0, 1.5)
        _indent(p, 24)
        _add_rich(p, line, 'FangSong', 'Times New Roman', 12)
        i += 1

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    sz = os.path.getsize(out_path)
    print(f'  -> {out_path} ({sz//1024} KB)')

if __name__ == '__main__':
    for md_p, out_p, fig_d in TARGETS:
        try:
            convert(md_p, out_p, fig_d)
        except Exception as e:
            print(f'ERROR: {md_p}: {e}')
            import traceback
            traceback.print_exc()
    print('Done.')
