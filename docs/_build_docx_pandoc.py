# -*- coding: utf-8 -*-
"""Build docx with pandoc OMML equations + python-docx images/tables."""
import re, os, subprocess, zipfile
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

BASE = os.path.dirname(os.path.abspath(__file__))
NS_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
NS_M = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"

TARGETS = [
    ("model_validation", "模型校验章节.md", "模型校验章节.docx"),
    ("simulation_design", "仿真设计说明.md", "仿真设计说明.docx"),
    ("program_design", "程序设计说明.md", "程序设计说明.docx"),
]

FIG_NAMES = {
    'Figure 1': 'fig1_overall_comparison.png',
    'Figure 2': 'fig2_per_round_timeseries.png',
    'Figure 3': 'fig3_error_structure.png',
    'Figure 4': 'fig4_f1_distribution.png',
    'Figure 5': 'fig5_cross_country_f1.png',
    'Figure 6': 'fig6_country_stability.png',
    'Figure 7': 'fig7_bloc_alignment.png',
    'Figure 8': 'fig8_italy_case_study.png',
    'Figure 9': 'fig9_gp_independence.png',
    'Figure 10': 'fig10_summary_dashboard.png',
    'Figure 11': 'fig1_f1_timeseries.png',
    'Figure 12': 'fig2_per_country.png',
    'Figure 13': 'fig3_confusion_matrix.png',
    'Figure 14': 'fig4_heatmap.png',
    'Figure 15': 'fig5_faction_evolution.png',
    'Figure 16': 'fig6_gp_independence.png',
}
OTHER_FIGS = {
    '三场景关键历史事件时间线': 'fig1_scene_timeline.png',
    '三场景国家数量与层级分布': 'fig3_country_distribution.png',
    '三场景追随格局演化对比': 'fig2_following_evolution.png',
    '国际体系结构演化谱系': 'fig4_system_evolution.png',
    '模型校验三场景时间线': 'fig4_three_scenes_timeline.png',
    'CINC综合国力指数计算架构': 'fig5_cinc_calculation.png',
    '四种领导类型行为偏好对比': 'fig3_leader_types_radar.png',
    '20项标准行为空间配置总览': 'fig2_action_space.png',
    '仿真单轮执行流程': 'fig1_simulation_flow.png',
    '核心仿真引擎模块架构': 'fig1_module_architecture.png',
    '数据持久化实体关系概览': 'fig2_data_er_diagram.png',
    '前端研究分析仪表板': 'fig3_frontend_dashboard.png',
    '前端并发控制机制': 'fig4_concurrency_control.png',
}
FIG_MAP = {**FIG_NAMES, **OTHER_FIGS}

def _font(run, cn, en='Times New Roman', size=12, bold=False, italic=False):
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None: rf = parse_xml(f'<w:rFonts {nsdecls("w")} />'); rPr.insert(0, rf)
    for k, v in [('eastAsia', cn), ('ascii', en), ('hAnsi', en), ('cs', en)]:
        rf.set(qn(f'w:{k}'), v)

def _sp(p, before=0, after=0, line=1.5):
    pf = p.paragraph_format; pf.space_before = Pt(before)
    pf.space_after = Pt(after); pf.line_spacing = line

def _indent(p, pt=24): p.paragraph_format.first_line_indent = Pt(pt)

def _add_rich(p, text, cn='FangSong', en='Times New Roman', size=12, bold=False):
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*)'; last = 0
    for m in re.finditer(pattern, text):
        if m.start() > last:
            r = p.add_run(text[last:m.start()]); _font(r, cn, en, size, bold)
        if m.group(2): r = p.add_run(m.group(2)); _font(r, cn, en, size, True)
        elif m.group(3): r = p.add_run(m.group(3)); _font(r, cn, en, size, bold, True)
        last = m.end()
    if last < len(text): r = p.add_run(text[last:]); _font(r, cn, en, size, bold)

def _is_table_sep(line): return bool(re.match(r'^\|[\s\-:|]+\|$', line.strip()))
def _is_table_row(line): s = line.strip(); return s.startswith('|') and s.endswith('|')
def _peek_table(lines, i):
    tbl = [lines[i]]; j = i + 1
    while j < len(lines):
        s = lines[j].strip()
        if _is_table_row(s) or _is_table_sep(s): tbl.append(lines[j]); j += 1
        else: break
    return tbl, j
def _add_table(doc, lines):
    clean = [l for l in lines if not _is_table_sep(l)]
    if len(clean) < 2: return
    header = [c.strip() for c in clean[0].split('|') if c.strip()]
    data = [[c.strip() for c in l.split('|') if c.strip()] for l in clean[1:] if l.strip()]
    data = [d for d in data if d]
    if not header: return
    ncols = len(header); nrows = 1 + len(data)
    table = doc.add_table(rows=nrows, cols=ncols); table.style = 'Table Grid'
    sz = 12 if ncols <= 5 else 9
    for ci, ct in enumerate(header):
        cell = table.cell(0, ci); cell.paragraphs[0].clear()
        _add_rich(cell.paragraphs[0], ct, 'SimHei', 'Times New Roman', sz)
        _sp(cell.paragraphs[0], 0, 0, 1.1)
    for ri, row in enumerate(data):
        for ci, ct in enumerate(row):
            if ci >= ncols: continue
            cell = table.cell(ri + 1, ci); cell.paragraphs[0].clear()
            _add_rich(cell.paragraphs[0], ct, 'FangSong', 'Times New Roman', sz)
            _sp(cell.paragraphs[0], 0, 0, 1.1)
    p = doc.add_paragraph(); _sp(p, 6, 0, 1.5)

def _tex_to_omml(latex_str):
    """Convert LaTeX equation to OMML paragraph element via pandoc."""
    latex_str = latex_str.strip()
    if not latex_str: return None
    tmp_md = os.path.join(BASE, '__eq_tmp.md')
    tmp_docx = os.path.join(BASE, '__eq_tmp.docx')
    try:
        with open(tmp_md, 'w', encoding='utf-8') as f:
            f.write(f"$$\n{latex_str}\n$$\n")
        subprocess.run(['pandoc', tmp_md, '-f', 'markdown+tex_math_dollars',
                        '-t', 'docx', '-o', tmp_docx],
                       capture_output=True, timeout=30, check=True)
        with zipfile.ZipFile(tmp_docx, 'r') as zf:
            doc_xml = etree.fromstring(zf.read('word/document.xml'))
        body = doc_xml.find(f'{NS_W}body')
        for p in body.findall(f'{NS_W}p'):
            if p.find(f'.//{NS_M}oMath') is not None or p.find(f'.//{NS_M}oMathPara') is not None:
                return p
        return None
    except Exception as e:
        print(f'  OMML warning: {e}')
        return None
    finally:
        for f in [tmp_md, tmp_docx]:
            if os.path.exists(f):
                try: os.remove(f)
                except: pass

def _image_path(alt, fig_dir):
    fname = FIG_MAP.get(alt, f"{alt}.png")
    fpath = os.path.join(fig_dir, fname)
    if os.path.exists(fpath): return fpath
    if alt.startswith('Figure '):
        n = alt.replace('Figure ', '')
        for f in sorted(os.listdir(fig_dir)):
            if f.startswith(f'fig{n}_'): return os.path.join(fig_dir, f)
    return None

def convert(md_path, out_path, fig_dir):
    print(f'Converting: {md_path}')
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')

    doc = Document()
    ns_el = doc.styles['Normal']
    for tag in (qn('w:pPr'), qn('w:rPr')):
        el = ns_el.element.find(tag)
        if el is not None: ns_el.element.remove(el)
    for sec in doc.sections:
        sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.18); sec.right_margin = Cm(3.18)

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line or line == '---':
            i += 1; continue

        # Image
        m_img = re.match(r'^!\[(.+?)\]\(.+\)$', line)
        if m_img:
            alt = m_img.group(1); fpath = _image_path(alt, fig_dir)
            if fpath:
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _sp(p, 3, 3, 1.5); r = p.add_run(); r.add_picture(fpath, width=Inches(5.5))
            if i + 1 < len(lines) and lines[i + 1].strip().startswith('*图：'):
                cap = re.sub(r'^\*(.+)\*$', r'\1', lines[i + 1].strip())
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _sp(p, 0, 3, 1.5); _add_rich(p, cap, 'FangSong', 'Times New Roman', 12)
                i += 1
            i += 1; continue

        # Table
        if _is_table_row(line):
            tbl, next_i = _peek_table(lines, i)
            print(f'  Table: {len(tbl)} rows')
            _add_table(doc, tbl); i = next_i; continue

        # Title (#)
        m1 = re.match(r'^# (.+)$', line)
        if m1 and not line.startswith('## '):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _sp(p, 12, 12, 2.0)
            _add_rich(p, m1.group(1), 'STXinwei', 'Times New Roman', 22); i += 1; continue

        # H1 (##)
        m2 = re.match(r'^## (.+)$', line)
        if m2:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _sp(p, 6, 0, 1.5)
            _add_rich(p, m2.group(1), 'SimHei', 'Times New Roman', 15, True); i += 1; continue

        # H2 (###)
        m3 = re.match(r'^### (.+)$', line)
        if m3:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _sp(p, 4, 0, 1.5)
            _add_rich(p, m3.group(1), 'KaiTi', 'Times New Roman', 14, True); i += 1; continue

        # H3 (####)
        m4 = re.match(r'^#### (.+)$', line)
        if m4:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _sp(p, 3, 0, 1.5)
            _add_rich(p, m4.group(1), 'FangSong', 'Times New Roman', 12, True); i += 1; continue

        # Math block $$...$$ (single-line)
        if line.startswith('$$') and line.endswith('$$') and len(line) > 4:
            latex = line[2:-2]
            omml_p = _tex_to_omml(latex)
            if omml_p is not None:
                body_el = doc.element.body
                p = doc.add_paragraph()
                body_el.replace(p._element, omml_p)
            else:
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _sp(p, 3, 3, 1.5); _add_rich(p, latex, 'Cambria Math', 'Cambria Math', 11)
            i += 1; continue

        # Math block $$...$$ (multi-line)
        if line == '$$':
            j = i + 1; latex_lines = []
            while j < len(lines):
                if lines[j].strip() == '$$': break
                latex_lines.append(lines[j]); j += 1
            if j < len(lines):
                omml_p = _tex_to_omml('\n'.join(latex_lines))
                if omml_p is not None:
                    body_el = doc.element.body
                    p = doc.add_paragraph()
                    body_el.replace(p._element, omml_p)
                else:
                    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _sp(p, 3, 3, 1.5)
                    _add_rich(p, '\n'.join(latex_lines), 'Cambria Math', 'Cambria Math', 11)
                i = j + 1; continue

        # Inline math $...$
        if '$' in line and not line.startswith('$$'):
            parts = re.split(r'(\$[^$]+\$)', line)
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _sp(p, 0, 0, 1.5); _indent(p, 24)
            for part in parts:
                if part.startswith('$') and part.endswith('$'):
                    omml_p = _tex_to_omml(part[1:-1])
                    if omml_p is not None:
                        for child in list(omml_p):
                            p._element.append(child)
                    else:
                        r = p.add_run(part[1:-1]); _font(r, 'Cambria Math', 'Cambria Math', 11)
                else:
                    _add_rich(p, part, 'FangSong', 'Times New Roman', 12)
            i += 1; continue

        # Bullet
        if re.match(r'^[-*]\s+', line):
            text = re.sub(r'^[-*]\s+', '', line)
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _sp(p, 0, 0, 1.5)
            pf = p.paragraph_format; pf.left_indent = Cm(1); pf.first_line_indent = Cm(-0.5)
            r = p.add_run('  '); _font(r, 'FangSong', 'Times New Roman', 12)
            _add_rich(p, text, 'FangSong', 'Times New Roman', 12); i += 1; continue

        # Blockquote
        if line.startswith('> '):
            text = line[2:]
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _sp(p, 0, 0, 1.5); pf = p.paragraph_format; pf.left_indent = Cm(1.5)
            _add_rich(p, text, 'KaiTi', 'Times New Roman', 12); i += 1; continue

        # Body
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _sp(p, 0, 0, 1.5); _indent(p, 24)
        _add_rich(p, line, 'FangSong', 'Times New Roman', 12); i += 1

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    sz = os.path.getsize(out_path)
    print(f'  -> {os.path.basename(out_path)} ({sz//1024} KB)')

def verify(docx_path):
    with zipfile.ZipFile(docx_path) as zf:
        tree = etree.fromstring(zf.read('word/document.xml'))
        eqs = len(tree.findall(f'.//{NS_M}oMath')) + len(tree.findall(f'.//{NS_M}oMathPara'))
        imgs = [f for f in zf.namelist() if f.startswith('word/media/image')]
        print(f'  Verify: equations={eqs}, images={len(imgs)}')

def main():
    for subdir, md_name, docx_name in TARGETS:
        md_path = os.path.join(BASE, subdir, md_name)
        out_path = os.path.join(BASE, subdir, docx_name)
        fig_dir = os.path.join(BASE, subdir, 'figures')
        if not os.path.exists(md_path):
            print(f'SKIP: {md_path}'); continue
        if os.path.exists(out_path):
            try: os.remove(out_path)
            except PermissionError: out_path = out_path.replace('.docx', '_NEW.docx')
        convert(md_path, out_path, fig_dir)
        verify(out_path)
    print('\nDone.')

if __name__ == '__main__':
    main()
