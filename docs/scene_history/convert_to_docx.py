# -*- coding: utf-8 -*-
"""Convert 三场景历史背景说明.md to Chinese academic DOCX."""
import re, os
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

MD_PATH = r"docs/scene_history/三场景历史背景说明.md"
FIG_DIR = r"docs/scene_history/figures"
OUT_PATH = r"docs/scene_history/三场景历史背景说明.docx"

FIG_MAP = {
    "三场景关键历史事件时间线": "fig1_scene_timeline.png",
    "三场景国家数量与层级分布": "fig3_country_distribution.png",
    "三场景追随格局演化对比": "fig2_following_evolution.png",
    "国际体系结构演化谱系": "fig4_system_evolution.png",
}


def _font(run, cn, en="Times New Roman", size=12, bold=False, italic=False):
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rPr.insert(0, rf)
    for k, v in [("eastAsia", cn), ("ascii", en), ("hAnsi", en), ("cs", en)]:
        rf.set(qn(f"w:{k}"), v)


def _spacing(para, before=0, after=0, line=1.5):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def _indent(para, pt=24):
    para.paragraph_format.first_line_indent = Pt(pt)


def _add_rich(para, text, cn="FangSong", en="Times New Roman", size=12, bold=False):
    """Add text with inline **bold** and *italic*."""
    pattern = r"(\*\*(.+?)\*\*|\*(.+?)\*)"
    last = 0
    for m in re.finditer(pattern, text):
        if m.start() > last:
            r = para.add_run(text[last : m.start()])
            _font(r, cn, en, size, bold)
        if m.group(2):
            r = para.add_run(m.group(2))
            _font(r, cn, en, size, True)
        elif m.group(3):
            r = para.add_run(m.group(3))
            _font(r, cn, en, size, bold, True)
        last = m.end()
    if last < len(text):
        r = para.add_run(text[last:])
        _font(r, cn, en, size, bold)


def _is_table_sep(line):
    return bool(re.match(r"^\|[\s\-:|]+\|$", line.strip()))


def _is_table_row(line):
    s = line.strip()
    return s.startswith("|") and s.endswith("|")


def _peek_table(lines, i):
    tbl = [lines[i]]
    j = i + 1
    while j < len(lines):
        s = lines[j].strip()
        if _is_table_row(s) or _is_table_sep(s):
            tbl.append(lines[j])
            j += 1
        else:
            break
    return tbl, j


def _add_table(doc, lines):
    clean = [l for l in lines if not _is_table_sep(l)]
    if len(clean) < 2:
        return
    header = [c.strip() for c in clean[0].split("|") if c.strip()]
    data = []
    for l in clean[1:]:
        cells = [c.strip() for c in l.split("|") if c.strip()]
        if cells:
            data.append(cells)
    if not header:
        return
    ncols = len(header)
    nrows = 1 + len(data)
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = "Table Grid"
    sz = 12 if ncols <= 5 else 9
    for ci, ct in enumerate(header):
        cell = table.cell(0, ci)
        cell.paragraphs[0].clear()
        _add_rich(cell.paragraphs[0], ct, "SimHei", "Times New Roman", sz)
        _spacing(cell.paragraphs[0], 0, 0, 1.1)
    for ri, row in enumerate(data):
        for ci, ct in enumerate(row):
            if ci >= ncols:
                continue
            cell = table.cell(ri + 1, ci)
            cell.paragraphs[0].clear()
            _add_rich(cell.paragraphs[0], ct, "FangSong", "Times New Roman", sz)
            _spacing(cell.paragraphs[0], 0, 0, 1.1)
    p = doc.add_paragraph()
    _spacing(p, 6, 0, 1.5)


def _is_image(line):
    return bool(re.match(r"^!\[(.+?)\]\(.+\)$", line.strip()))


def _image_alt(line):
    m = re.match(r"^!\[(.+?)\]\(.+\)$", line.strip())
    return m.group(1) if m else None


def main():
    print("Reading markdown...")
    with open(MD_PATH, "r", encoding="utf-8") as f:
        content = f.read()
    lines = content.split("\n")

    doc = Document()
    ns_el = doc.styles["Normal"]
    for tag in (qn("w:pPr"), qn("w:rPr")):
        el = ns_el.element.find(tag)
        if el is not None:
            ns_el.element.remove(el)

    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.18)
        sec.right_margin = Cm(3.18)

    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue
        if line == "---":
            i += 1
            continue

        # Image
        if _is_image(line):
            alt = _image_alt(line)
            fname = FIG_MAP.get(alt, alt)
            fpath = os.path.join(FIG_DIR, fname)
            if os.path.exists(fpath):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _spacing(p, 3, 3, 1.5)
                r = p.add_run()
                r.add_picture(fpath, width=Inches(5.5))
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _spacing(p, 3, 3, 1.5)
                r = p.add_run(f"[图：{alt}]")
                _font(r, "FangSong", "Times New Roman", 12)
            # image caption
            if i + 1 < len(lines) and lines[i + 1].strip().startswith("*图："):
                cap = re.sub(r"^\*(.+)\*$", r"\1", lines[i + 1].strip())
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _spacing(p, 0, 3, 1.5)
                _add_rich(p, cap, "FangSong", "Times New Roman", 12)
                i += 1
            i += 1
            continue

        # Table
        if _is_table_row(line):
            tbl, next_i = _peek_table(lines, i)
            print(f"  Table: {len(tbl)} rows")
            _add_table(doc, tbl)
            i = next_i
            continue

        # Title (# )
        m1 = re.match(r"^# (.+)$", line)
        if m1 and not line.startswith("## "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _spacing(p, 12, 12, 2.0)
            _add_rich(p, m1.group(1), "STSong", "Times New Roman", 22)
            i += 1
            continue

        # H1 (## )
        m2 = re.match(r"^## (.+)$", line)
        if m2:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _spacing(p, 6, 0, 1.5)
            _add_rich(p, m2.group(1), "SimHei", "Times New Roman", 15, True)
            i += 1
            continue

        # H2 (### )
        m3 = re.match(r"^### (.+)$", line)
        if m3:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _spacing(p, 4, 0, 1.5)
            _add_rich(p, m3.group(1), "KaiTi", "Times New Roman", 14, True)
            i += 1
            continue

        # H3 (#### )
        m4 = re.match(r"^#### (.+)$", line)
        if m4:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _spacing(p, 3, 0, 1.5)
            _add_rich(p, m4.group(1), "FangSong", "Times New Roman", 12, True)
            i += 1
            continue

        # Bullet
        if re.match(r"^[-*]\s+", line):
            text = re.sub(r"^[-*]\s+", "", line)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _spacing(p, 0, 0, 1.5)
            pf = p.paragraph_format
            pf.left_indent = Cm(1)
            pf.first_line_indent = Cm(-0.5)
            r = p.add_run("• ")
            _font(r, "FangSong", "Times New Roman", 12)
            _add_rich(p, text, "FangSong", "Times New Roman", 12)
            i += 1
            continue

        # Body
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _spacing(p, 0, 0, 1.5)
        _indent(p, 24)
        _add_rich(p, line, "FangSong", "Times New Roman", 12)
        i += 1

    print(f"Saving to {OUT_PATH}...")
    doc.save(OUT_PATH)
    print("Done!")


if __name__ == "__main__":
    main()
